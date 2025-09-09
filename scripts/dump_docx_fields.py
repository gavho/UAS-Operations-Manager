import sys, os, re
from typing import Dict, List, Tuple

try:
    from docx import Document
except Exception as e:
    print("python-docx is required: pip install python-docx")
    sys.exit(2)


def dump_merge_fields(doc) -> Dict[str, str]:
    values: Dict[str, str] = {}

    def all_paragraphs():
        pars = list(getattr(doc, 'paragraphs', []) or [])
        # Tables in body
        for tbl in getattr(doc, 'tables', []) or []:
            for row in getattr(tbl, 'rows', []) or []:
                for cell in getattr(row, 'cells', []) or []:
                    pars.extend(getattr(cell, 'paragraphs', []) or [])
                    for nt in getattr(cell, 'tables', []) or []:
                        for nrow in getattr(nt, 'rows', []) or []:
                            for ncell in getattr(nrow, 'cells', []) or []:
                                pars.extend(getattr(ncell, 'paragraphs', []) or [])
        # Headers/Footers
        for sec in getattr(doc, 'sections', []) or []:
            for hf in [sec.header, sec.footer]:
                if hf is None:
                    continue
                pars.extend(getattr(hf, 'paragraphs', []) or [])
                for tbl in getattr(hf, 'tables', []) or []:
                    for row in getattr(tbl, 'rows', []) or []:
                        for cell in getattr(row, 'cells', []) or []:
                            pars.extend(getattr(cell, 'paragraphs', []) or [])
        return pars

    # Complex fields (fldChar/instrText)
    def process_paragraph_complex_fields(p) -> None:
        in_field = False
        field_name = None
        capturing_value = False
        value_runs: List[str] = []
        for run in getattr(p, 'runs', []) or []:
            r = run._r
            fldChar = r.find('.//w:fldChar', r.nsmap)
            instrText = r.find('.//w:instrText', r.nsmap)
            if fldChar is not None and fldChar.get('{%s}fldCharType' % fldChar.nsmap['w']) == 'begin':
                in_field, field_name, capturing_value, value_runs = True, None, False, []
                continue
            if in_field and instrText is not None:
                m = re.search(r'MERGEFIELD\s+([\w\-.]+)', (instrText.text or ''), re.I)
                if m:
                    field_name = m.group(1).strip()
                continue
            if in_field and fldChar is not None and fldChar.get('{%s}fldCharType' % fldChar.nsmap['w']) == 'separate':
                capturing_value = True
                value_runs = []
                continue
            if in_field and fldChar is not None and fldChar.get('{%s}fldCharType' % fldChar.nsmap['w']) == 'end':
                if field_name:
                    values[field_name.lower()] = ''.join(value_runs).strip()
                in_field, field_name, capturing_value, value_runs = False, None, False, []
                continue
            if in_field and capturing_value:
                for t in r.findall('.//w:t', r.nsmap):
                    if t.text:
                        value_runs.append(t.text)

    for p in getattr(doc, 'paragraphs', []) or []:
        process_paragraph_complex_fields(p)
    for tbl in getattr(doc, 'tables', []) or []:
        for row in getattr(tbl, 'rows', []) or []:
            for cell in getattr(row, 'cells', []) or []:
                for p in getattr(cell, 'paragraphs', []) or []:
                    process_paragraph_complex_fields(p)
    for sec in getattr(doc, 'sections', []) or []:
        for hf in [sec.header, sec.footer]:
            if hf is None:
                continue
            for p in getattr(hf, 'paragraphs', []) or []:
                process_paragraph_complex_fields(p)
            for tbl in getattr(hf, 'tables', []) or []:
                for row in getattr(tbl, 'rows', []) or []:
                    for cell in getattr(row, 'cells', []) or []:
                        for p in getattr(cell, 'paragraphs', []) or []:
                            process_paragraph_complex_fields(p)

    # Simple fields (fldSimple)
    for p in all_paragraphs():
        el = p._p
        for fs in el.findall('.//w:fldSimple', el.nsmap):
            instr = fs.get('{%s}instr' % fs.nsmap['w'], '')
            m = re.search(r'MERGEFIELD\s+([\w\-.]+)', instr, re.I)
            if m:
                name = m.group(1).strip().lower()
                ts = fs.findall('.//w:t', fs.nsmap)
                value = ''.join([t.text or '' for t in ts]).strip()
                values[name] = value
    return values


def dump_content_controls(doc) -> List[Tuple[str, str, str]]:
    out: List[Tuple[str, str, str]] = []
    # Body
    root = doc.part.element.body
    for sdt in root.findall('.//w:sdt', root.nsmap):
        alias = sdt.find('.//w:alias', root.nsmap)
        tag = sdt.find('.//w:tag', root.nsmap)
        title = alias.get('{%s}val' % alias.nsmap['w']) if alias is not None else ''
        tagval = tag.get('{%s}val' % tag.nsmap['w']) if tag is not None else ''
        txt = ''.join([(t.text or '') for t in sdt.findall('.//w:t', root.nsmap)]).strip()
        out.append((title, tagval, txt))
    # Headers/Footers
    for sec in getattr(doc, 'sections', []) or []:
        for part in [sec.header.part.element, sec.footer.part.element]:
            if part is None:
                continue
            for sdt in part.findall('.//w:sdt', part.nsmap):
                alias = sdt.find('.//w:alias', part.nsmap)
                tag = sdt.find('.//w:tag', part.nsmap)
                title = alias.get('{%s}val' % alias.nsmap['w']) if alias is not None else ''
                tagval = tag.get('{%s}val' % tag.nsmap['w']) if tag is not None else ''
                txt = ''.join([(t.text or '') for t in sdt.findall('.//w:t', part.nsmap)]).strip()
                out.append((title, tagval, txt))
    return out


def dump_plain_text(doc) -> Tuple[List[str], List[str]]:
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    rows = []
    for tbl in doc.tables:
        for row in tbl.rows:
            rows.append(' | '.join(c.text.strip() for c in row.cells))
    return paras, rows


def main(path: str):
    if not os.path.isfile(path):
        print(f"File not found: {path}")
        sys.exit(1)
    doc = Document(path)
    print('== MERGEFIELDS ==')
    mf = dump_merge_fields(doc)
    if not mf:
        print('(none detected)')
    else:
        for k in sorted(mf.keys()):
            print(f'{k}: {mf[k]}')

    print('\n== CONTENT CONTROLS (w:sdt) ==')
    sdt = dump_content_controls(doc)
    if not sdt:
        print('(none detected)')
    else:
        for title, tag, txt in sdt:
            print(f'Alias: {title} | Tag: {tag} | Text: {txt}')

    print('\n== TEXT (paragraphs) ==')
    paras, rows = dump_plain_text(doc)
    for p in paras[:100]:
        print(p)

    print('\n== TABLE ROWS ==')
    for r in rows[:100]:
        print(r)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python scripts/dump_docx_fields.py <docx_path>')
        sys.exit(1)
    main(sys.argv[1])
